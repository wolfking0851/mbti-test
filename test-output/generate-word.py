#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试文件生成 - Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_test_word():
    doc = Document()
    
    # 标题
    title = doc.add_heading('🦞 智能龙虾 - Word 测试文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph('能力验证测试文件')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    doc.add_paragraph()
    
    # 基本信息
    doc.add_heading('1️⃣ 基本信息', level=1)
    doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'生成者：大娃 (SF-0001)')
    doc.add_paragraph(f'测试地点：贵阳')
    
    # 功能列表
    doc.add_heading('2️⃣ 支持的 Word 功能', level=1)
    
    # 无序列表
    doc.add_paragraph('基础功能：', style='List Bullet')
    doc.add_paragraph('✅ 标题和段落', style='List Bullet')
    doc.add_paragraph('✅ 无序列表', style='List Bullet')
    doc.add_paragraph('✅ 有序列表', style='List Bullet')
    doc.add_paragraph('✅ 文字格式（粗体、斜体、颜色）', style='List Bullet')
    
    # 有序列表
    doc.add_paragraph('高级功能：', style='List Number')
    doc.add_paragraph('表格支持', style='List Number')
    doc.add_paragraph('样式定制', style='List Number')
    doc.add_paragraph('分页控制', style='List Number')
    
    # 表格
    doc.add_heading('3️⃣ 表格示例', level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['功能', '状态', '备注']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # 数据行
    data = [
        ['文档生成', '✅ 正常', 'python-docx'],
        ['表格支持', '✅ 正常', '3 列 4 行'],
        ['中文支持', '✅ 正常', '无乱码']
    ]
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data
    
    doc.add_paragraph()
    
    # 颜色示例
    doc.add_heading('4️⃣ 颜色示例', level=1)
    p = doc.add_paragraph()
    p.add_run('红色文字 ').bold = True
    p.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    
    p.add_run('绿色文字 ').bold = True
    p.runs[1].font.color.rgb = RGBColor(0, 128, 0)
    
    p.add_run('蓝色文字').bold = True
    p.runs[2].font.color.rgb = RGBColor(0, 0, 255)
    
    # 结束
    doc.add_paragraph()
    doc.add_paragraph('—' * 30)
    doc.add_paragraph('测试完成！✅', style='Intense Quote')
    
    # 保存
    output_path = '/home/admin/.openclaw/workspace/test-output/测试文档-Word.docx'
    doc.save(output_path)
    print(f'✅ Word 文档已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    create_test_word()
